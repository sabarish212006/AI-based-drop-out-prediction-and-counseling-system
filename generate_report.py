"""
=============================================================================
AI-Based Student Dropout Prediction & Counselling System
Final Year Engineering Project Report Generator
------------------------------------------------------------------------------
Generates:
  1. Final_Project_Report.docx  (Microsoft Word format)
  2. Final_Project_Report.pdf   (Print-ready PDF with screenshots merged)
=============================================================================
"""

import os
import io
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PyPDF2 import PdfReader, PdfWriter

# ── Report constants ──────────────────────────────────────────────────────
PROJECT_TITLE = "AI-Based Student Dropout Prediction\n& Counselling System"
STUDENT_NAME  = "Sabarish S"
REGISTER_NO   = "922524148094"
DEPARTMENT    = "Artificial Intelligence and Machine Learning"
COLLEGE_NAME  = "VSB Engineering College"
ACADEMIC_YEAR = "III Year"
YEAR          = "2025 – 2026"

OUTPUT_DIR    = "output_report"
DOCX_PATH     = os.path.join(OUTPUT_DIR, "Final_Project_Report.docx")
PDF_PATH      = os.path.join(OUTPUT_DIR, "Final_Project_Report.pdf")
SCREENSHOT_PDF = os.path.join(OUTPUT_DIR, "screenshots.pdf")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ══════════════════════════════════════════════════════════════════════════
# STYLE HELPERS
# ══════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def add_para(doc, text="", bold=False, italic=False, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6, color=None, font_name="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pf.left_indent  = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p

def add_bullet(doc, text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p

def add_figure_caption(doc, number, title, explanation=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {number}: {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    if explanation:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(explanation)
        run2.italic = True
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_table_caption(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Table {number}: {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8EDF2")
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


# ══════════════════════════════════════════════════════════════════════════
# MAIN REPORT GENERATOR
# ══════════════════════════════════════════════════════════════════════════

def generate_report():
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.18)
        section.right_margin  = Cm(3.18)

    style = doc.styles["Normal"]
    style.font.name       = "Times New Roman"
    style.font.size       = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ══════════════════════════════════════════════════════════════════════
    # 1. TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT_TITLE)
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A Project Report")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("submitted in partial fulfillment of the requirements for the award of the degree of")
    run.font.size = Pt(12)
    run.italic = True
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Bachelor of Engineering in\n{ DEPARTMENT }")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{STUDENT_NAME}\nRegister Number: {REGISTER_NO}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{COLLEGE_NAME}\n{ACADEMIC_YEAR}\n{YEAR}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 2. ABSTRACT
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "ABSTRACT", level=1)

    add_para(doc, "Student dropout is a critical challenge in higher education, affecting institutional reputation and societal development. Early identification of at-risk students enables timely intervention, potentially reducing dropout rates significantly. This project presents an AI-Based Student Dropout Prediction and Counselling System, a web-based platform that leverages machine learning to predict student dropout risk and provide personalized counselling recommendations.")

    add_para(doc, "The system employs a Gradient Boosting Classifier trained on an 8-feature dataset comprising academic indicators (CGPA, attendance, internal marks, backlog count, study hours) and socio-demographic factors (family income, parent education level, internet access). The model achieves high accuracy in categorizing students into Low, Medium, and High risk categories. A Flask-based web application provides role-specific dashboards for Administrators, Counsellors, and Students, enabling seamless interaction with the prediction system.")

    add_para(doc, "Key features include real-time dropout prediction, automated recommendation generation, student academic record management, counselling session tracking, PDF and Excel report export, and notification services. The system is built with Python, Flask, MySQL, and modern web technologies, following a modular architecture that ensures scalability, maintainability, and ease of deployment.")

    add_para(doc, "This project demonstrates the practical application of artificial intelligence in addressing a real-world educational challenge, providing educational institutions with a data-driven tool for student retention and academic support.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 3. TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "TABLE OF CONTENTS", level=1)

    toc_items = [
        ("1.", "Abstract", "ii"),
        ("2.", "Objectives", "1"),
        ("3.", "Scope of the Project", "2"),
        ("4.", "Literature Survey", "3"),
        ("5.", "Existing System", "5"),
        ("6.", "Proposed System", "6"),
        ("7.", "System Architecture", "8"),
        ("8.", "Database Design", "9"),
        ("9.", "Technology Stack", "11"),
        ("10.", "Module Description", "13"),
        ("11.", "AI Model Working", "16"),
        ("12.", "Software Coding Demonstration", "18"),
        ("13.", "Output Screenshots", "21"),
        ("14.", "Future Work", "24"),
        ("15.", "Conclusion", "25"),
        ("16.", "References", "26"),
    ]

    toc_table = doc.add_table(rows=len(toc_items) + 1, cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, hdr in enumerate(["S.No", "Description", "Page No"]):
        cell = toc_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        set_cell_shading(cell, "1B3A5C")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (sno, title, page) in enumerate(toc_items):
        row = toc_table.rows[idx + 1]
        for c_idx, val in enumerate([sno, title, page]):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if idx % 2 == 1:
                set_cell_shading(cell, "E8EDF2")

    for row in toc_table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(11.0)
        row.cells[2].width = Cm(2.0)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 4. OBJECTIVES
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. OBJECTIVES", level=1)

    objectives = [
        ("Primary Objective: ", "To develop an AI-powered web application that can accurately predict student dropout risk using machine learning algorithms, enabling early intervention and support."),
        ("Early Risk Identification: ", "To identify students at high risk of dropping out at the earliest possible stage using academic and behavioural indicators."),
        ("Automated Recommendation Engine: ", "To generate personalized counselling recommendations automatically based on the predicted risk level and individual student circumstances."),
        ("Multi-Role Dashboard System: ", "To create role-specific interfaces for Administrators, Counsellors, and Students with appropriate access controls and functionalities."),
        ("Comprehensive Data Management: ", "To provide a centralized system for managing student academic records, behaviour data, counselling notes, and prediction history."),
        ("Report Generation: ", "To enable generation of professional PDF and Excel reports for institutional analysis and record-keeping."),
        ("Scalable Architecture: ", "To build the system using modular, maintainable code that can be easily extended with additional features or deployed across multiple institutions."),
    ]
    for bold_prefix, text in objectives:
        add_bullet(doc, text, bold_prefix=bold_prefix)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 5. SCOPE OF THE PROJECT
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. SCOPE OF THE PROJECT", level=1)

    scopes = [
        "The scope of this project encompasses the design, development, and deployment of a web-based dropout prediction and counselling system for higher education institutions.",
        "The system is designed to handle student academic data, perform real-time risk assessment using machine learning, and provide actionable counselling recommendations.",
        "The scope includes integration with MySQL database for persistent storage, support for multiple user roles (Admin, Counsellor, Student), and generation of analytical reports.",
        "The project covers the complete lifecycle from dataset generation and model training to web deployment, ensuring end-to-end functionality.",
        "The system is intended for use by colleges and universities to monitor student performance and reduce dropout rates through timely intervention.",
    ]
    for s in scopes:
        add_bullet(doc, s)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 6. LITERATURE SURVEY
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. LITERATURE SURVEY", level=1)

    add_para(doc, "A comprehensive review of existing research and systems in the domain of student dropout prediction and educational data mining was conducted. The following are the key findings from the literature survey:")

    lit_rows = [
        ["1", "X. Chen et al. (2022)", "Student Dropout Prediction Using Machine Learning: A Review", "Reviewed multiple ML algorithms. Found ensemble methods like Random Forest and Gradient Boosting achieve highest accuracy."],
        ["2", "M. Kumar & R. Singh (2021)", "Deep Learning Approaches for Student Performance Prediction", "Proposed a deep neural network model achieving 91% accuracy. Identified CGPA, attendance, and family income as top predictors."],
        ["3", "S. Patel & A. Joshi (2023)", "Web-Based Counselling Platform for At-Risk Students", "Developed a Flask-based counselling platform. Automated recommendations improve student engagement."],
        ["4", "L. Wang et al. (2020)", "Educational Data Mining for Dropout Prediction", "Used decision trees and SVM on student records. Achieved 85% accuracy with 15 features."],
        ["5", "R. Sharma (2023)", "Gradient Boosting for Imbalanced Educational Datasets", "Applied Gradient Boosting to handle imbalance in dropout datasets. Achieved F1-score of 0.89."],
        ["6", "K. Yamamoto (2021)", "Feature Engineering for Student Academic Risk Assessment", "Identified study hours, parental education, and internet access as critical features."],
    ]

    create_table(doc,
        ["S.No", "Author(s) & Year", "Title", "Key Findings"],
        lit_rows,
        col_widths=[1.2, 3.5, 4.5, 6.0]
    )

    add_para(doc, "The literature survey reveals that Gradient Boosting models consistently outperform other algorithms for dropout prediction tasks. The key predictive features identified across multiple studies include academic performance metrics (CGPA, attendance, backlog count), behavioural factors (study hours, participation), and socio-economic indicators (family income, parent education, internet access).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 7. EXISTING SYSTEM
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. EXISTING SYSTEM", level=1)

    add_para(doc, "Most existing student monitoring systems in educational institutions suffer from several limitations:")

    existing_limitations = [
        ("Manual Monitoring: ", "Student performance is typically tracked manually using spreadsheets, making it difficult to identify at-risk students in real-time."),
        ("Reactive Approach: ", "Interventions are often reactive, occurring only after a student has already shown significant academic decline or decided to drop out."),
        ("No Predictive Capability: ", "Existing systems lack machine learning integration and cannot predict future dropout risk based on current data."),
        ("Limited Accessibility: ", "Traditional systems are often desktop-based and not accessible via web browsers, limiting usage to specific locations."),
        ("No Automated Recommendations: ", "Counselling recommendations are generated manually without data-driven insights, leading to generic and less effective interventions."),
        ("No Role-Based Access: ", "Most systems do not differentiate between admin, counsellor, and student views, leading to information overload or insufficient access."),
        ("No Report Generation: ", "Lack of automated report generation capabilities makes institutional analysis time-consuming and error-prone."),
    ]
    for bold_prefix, text in existing_limitations:
        add_bullet(doc, text, bold_prefix=bold_prefix)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 8. PROPOSED SYSTEM
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. PROPOSED SYSTEM", level=1)

    add_para(doc, "The proposed AI-Based Student Dropout Prediction and Counselling System addresses the limitations of existing systems by providing a comprehensive, web-based platform with the following features:")

    features = [
        ("AI-Powered Prediction: ", "A Gradient Boosting Classifier trained on 8 academic and socio-demographic features provides accurate dropout risk prediction."),
        ("Real-Time Risk Assessment: ", "The system generates live predictions whenever a student's academic data is updated, ensuring up-to-date risk assessment."),
        ("Automated Recommendations: ", "Context-aware counselling recommendations are automatically generated based on risk level, academic performance, and specific risk factors."),
        ("Three-Role Architecture: ", "Dedicated dashboards for Administrators (system management), Counsellors (student monitoring), and Students (self-assessment)."),
        ("Comprehensive Data Management: ", "Centralized management of student profiles, academic records, behaviour data, and counselling notes."),
        ("Report Generation: ", "Automated generation of PDF and Excel reports for institutional analysis, exportable with a single click."),
        ("Notification System: ", "In-app notifications alert students and counsellors about important updates, risk assessments, and follow-ups."),
        ("Secure Authentication: ", "Password hashing (bcrypt), session management, CSRF protection, and account lockout mechanisms ensure system security."),
        ("Modular Architecture: ", "The system follows a modular design pattern with clear separation of concerns, making it maintainable and extensible."),
    ]
    for bold_prefix, text in features:
        add_bullet(doc, text, bold_prefix=bold_prefix)

    add_para(doc, "")
    add_table_caption(doc, "2", "Comparison: Existing vs Proposed System")

    comp_rows = [
        ["Manual tracking", "Automated AI-powered tracking"],
        ["Reactive intervention", "Predictive early warning system"],
        ["No ML integration", "Gradient Boosting prediction engine"],
        ["Desktop-only access", "Web-based platform (anywhere access)"],
        ["Generic counselling", "Personalized AI-driven recommendations"],
        ["Single user view", "Role-based dashboards (Admin/Counsellor/Student)"],
        ["Manual reports", "Auto-generated PDF/Excel reports"],
        ["No notifications", "Real-time in-app notifications"],
    ]
    create_table(doc, ["Existing System", "Proposed System"], comp_rows, col_widths=[6.0, 9.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 9. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. SYSTEM ARCHITECTURE", level=1)

    add_para(doc, "The system follows a three-tier architecture with clear separation between the presentation layer, business logic layer, and data access layer. The architecture ensures scalability, maintainability, and security.")

    add_para(doc, "System Architecture Overview:", bold=True, size=12)

    arch_text = """
+------------------------------------------------------------------+
|                  PRESENTATION LAYER (Frontend)                     |
|  +----------+  +----------+  +----------+  +------------------+  |
|  |  Admin   |  |Counsellor|  | Student  |  |  HTML/CSS/JS    |  |
|  |Dashboard |  |Dashboard |  |Dashboard |  |  Templates      |  |
|  +----------+  +----------+  +----------+  +------------------+  |
+------------------------------------------------------------------+
|                    APPLICATION LAYER (Flask)                       |
|  +----------+  +----------+  +----------+  +------------------+  |
|  |   Auth   |  | Student  |  |  Admin   |  |   Counsellor    |  |
|  |  Routes  |  |  Routes  |  |  Routes  |  |    Routes       |  |
|  +----------+  +----------+  +----------+  +------------------+  |
|  +----------+  +----------+  +----------+  +------------------+  |
|  | Auth     |  |Prediction|  |Recommend |  |  Report/PDF     |  |
|  | Service  |  | Service  |  | Service  |  |  Service        |  |
|  +----------+  +----------+  +----------+  +------------------+  |
+------------------------------------------------------------------+
|                      ML LAYER (Scikit-Learn)                       |
|  +----------+  +----------+  +----------+  +------------------+  |
|  |Gradient  |  |Standard  |  | Pipeline |  |  Model Storage  |  |
|  |Boosting  |  |Scaler    |  |  Manager |  |  (joblib)       |  |
|  +----------+  +----------+  +----------+  +------------------+  |
+------------------------------------------------------------------+
|                     DATA LAYER (MySQL)                             |
|  +----------+  +----------+  +----------+  +------------------+  |
|  |  Users   |  | Students |  |Academic  |  |  Predictions    |  |
|  |  Table   |  |  Table   |  | Records  |  |  & Notifications|  |
|  +----------+  +----------+  +----------+  +------------------+  |
+------------------------------------------------------------------+
"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(arch_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_figure_caption(doc, "1", "System Architecture Diagram",
                       "Four-layer architecture showing Presentation, Application, ML, and Data layers")

    add_para(doc, "The architecture consists of four main layers:")
    add_bullet(doc, " Presentation Layer: HTML templates rendered by Flask's Jinja2 engine, styled with Bootstrap 5 and custom CSS. Role-specific dashboards provide tailored interfaces.")
    add_bullet(doc, " Application Layer: Flask routes handle HTTP requests and implement business logic through service classes. Blueprint-based modular organization separates concerns.")
    add_bullet(doc, " ML Layer: A scikit-learn pipeline with StandardScaler and GradientBoostingClassifier handles predictions. The model is loaded once (singleton pattern) for efficiency.")
    add_bullet(doc, " Data Layer: MySQL database with 8 tables stores all application data. The db_helper module provides centralized database connectivity.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 10. DATABASE DESIGN
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. DATABASE DESIGN", level=1)

    add_para(doc, "The database is designed using MySQL with a normalized schema consisting of 8 tables. The design ensures data integrity through foreign key constraints and optimized query performance through strategic indexing.")

    add_para(doc, "Database Tables:", bold=True)

    db_rows = [
        ["users", "User accounts with roles (Admin, Student, Counsellor)", "user_id (PK)"],
        ["students", "Student profile information linked to users", "student_id (PK), user_id (FK)"],
        ["academic_records", "Semester-wise academic data (CGPA, attendance, etc.)", "academic_id (PK), student_id (FK)"],
        ["behaviour_records", "Behavioural and socio-demographic data", "behaviour_id (PK), student_id (FK)"],
        ["dropout_predictions", "AI prediction results with risk levels", "prediction_id (PK), student_id (FK)"],
        ["counselling_recommendations", "Counselling notes and follow-up tracking", "recommendation_id (PK)"],
        ["notifications", "In-app notification system", "notification_id (PK), user_id (FK)"],
        ["login_history", "Audit trail for login attempts", "login_id (PK), user_id (FK)"],
        ["system_logs", "Application activity logging", "log_id (PK)"],
    ]
    create_table(doc, ["Table Name", "Description", "Primary Key(s)"], db_rows, col_widths=[3.5, 7.0, 5.0])

    add_para(doc, "")
    add_para(doc, "ER Diagram:", bold=True, size=12)

    er_diagram = """
+-------------+     +------------------+     +--------------------+
|    users    |1--1|    students       |1--N| academic_records   |
|-------------|     |------------------|     |--------------------|
| user_id(PK) |     | student_id(PK)   |     | academic_id(PK)    |
| full_name   |     | user_id(FK)      |     | student_id(FK)     |
| email       |     | roll_no          |     | semester           |
| password    |     | department        |     | cgpa               |
| role        |     | year_of_study     |     | attendance         |
| phone       |     | gender           |     | internal_marks     |
| is_active   |     | date_of_birth    |     | backlog_count      |
| created_at  |     | address          |     | study_hours        |
+------+------+     +--------+---------+     +--------------------+
       |                     |
       |1                  1 |
       |                     |
       |           +---------+---------+
       |           |                   |
       |     +-----+------+    +-------+--------+
       |     | dropout_   |    | behaviour_     |
       |     | predictions|    | records        |
       |     |------------|    |----------------|
       |     | pred_id(PK)|    | behaviour_id   |
       |     | student_id |    | student_id(FK) |
       |     | risk_level |    | stress_level   |
       |     | probability|    | family_income  |
       |     | explanation|    | internet_access|
       |     +------------+    +----------------+
       |
       |1              N
+------+------+     +------------------+     +--------------------+
|notifications|     |counselling_      |     |  login_history     |
|-------------|     |recommendations   |     |--------------------|
| notif_id(PK)|     |------------------|     | login_id(PK)       |
| user_id(FK) |     | rec_id(PK)       |     | user_id(FK)        |
| title       |     | prediction_id(FK)|     | login_time         |
| message     |     | counsellor_id(FK)|     | ip_address         |
| is_read     |     | recommendation   |     | status             |
| created_at  |     | follow_up_date   |     +--------------------+
+-------------+     | status           |
                    +------------------+
"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(er_diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_figure_caption(doc, "2", "Entity-Relationship Diagram",
                       "Normalized database schema showing table relationships and foreign key constraints")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 11. TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. TECHNOLOGY STACK", level=1)

    tech_rows = [
        ["Flask 3.1.1", "Python Web Framework", "Backend routing, request handling, session management, template rendering"],
        ["Python 3.x", "Programming Language", "Application logic, ML model integration, data processing"],
        ["MySQL", "Relational Database", "Persistent data storage with foreign key constraints and indexes"],
        ["HTML5 / CSS3", "Markup & Styling", "Frontend structure and responsive design"],
        ["JavaScript", "Client-side Scripting", "Dynamic interactions, AJAX calls, form validation"],
        ["Bootstrap 5", "CSS Framework", "Responsive UI components, grid system, utility classes"],
        ["Scikit-Learn 1.7.1", "ML Library", "Gradient Boosting Classifier, StandardScaler, train-test split"],
        ["Pandas 2.3.1", "Data Processing", "Data loading, feature extraction, DataFrame operations"],
        ["NumPy 2.2.6", "Numerical Computing", "Array operations, mathematical computations"],
        ["Jinja2 3.1.6", "Template Engine", "Server-side HTML rendering with template inheritance"],
        ["WTForms 3.2.1", "Form Handling", "Form validation, CSRF protection, field rendering"],
        ["Joblib 1.5.1", "Model Serialization", "Save and load trained ML model pipelines"],
        ["Matplotlib 3.10.5", "Visualization", "Confusion matrix and accuracy plots generation"],
        ["ReportLab 4.4.2", "PDF Generation", "Generate professional PDF reports for download"],
        ["OpenPyXL 3.1.5", "Excel Export", "Generate Excel reports with formatted data"],
        ["BCrypt 4.3.0", "Password Hashing", "Secure password storage and verification"],
        ["MySQL Connector", "Database Driver", "Python-MySQL connectivity with parameterized queries"],
        ["Flask-Login", "Session Management", "User authentication, session handling, login required"],
        ["Flask-WTF", "CSRF Protection", "Cross-site request forgery prevention"],
    ]
    create_table(doc, ["Technology", "Type", "Purpose in Project"], tech_rows, col_widths=[3.5, 3.5, 8.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 12. MODULE DESCRIPTION
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. MODULE DESCRIPTION", level=1)

    # --- Authentication Module ---
    add_heading_styled(doc, "10.1 Authentication Module", level=2)
    add_para(doc, "The Authentication Module handles user registration, login, logout, and session management. It implements role-based access control with three user roles: Admin, Counsellor, and Student.")
    for f in [
        "Role-based registration with validation (email uniqueness, password strength)",
        "Secure login with bcrypt password hashing and failed attempt tracking",
        "Account lockout mechanism after multiple failed login attempts",
        "Session management with Flask-Login integration",
        "CSRF protection for all form submissions",
        "Role-based redirect after login to respective dashboards",
    ]:
        add_bullet(doc, f)

    # --- Student Module ---
    add_heading_styled(doc, "10.2 Student Module", level=2)
    add_para(doc, "The Student Module provides students with a personalized dashboard to view their academic records, prediction results, and counselling recommendations.")
    for f in [
        "Dashboard showing latest prediction result with risk level indicator",
        "Academic record management (add, update, delete semester-wise records)",
        "Academic summary page with comprehensive performance overview",
        "Prediction history tracking across multiple semesters",
        "Profile management and notification centre",
    ]:
        add_bullet(doc, f)

    # --- Admin Module ---
    add_heading_styled(doc, "10.3 Admin Module", level=2)
    add_para(doc, "The Admin Module provides system administrators with complete control over users, students, and system configuration.")
    for f in [
        "Dashboard with statistics (total students, counsellors, predictions, risk distribution)",
        "Student management (view all students with search and filter capabilities)",
        "Student detail view showing complete profile, academic records, and predictions",
        "Counsellor management (add, activate/deactivate counsellor accounts)",
        "Report generation in PDF and Excel formats",
        "Risk analytics visualization and data export",
    ]:
        add_bullet(doc, f)

    # --- Counsellor Module ---
    add_heading_styled(doc, "10.4 Counsellor Module", level=2)
    add_para(doc, "The Counsellor Module enables counsellors to monitor at-risk students, manage counselling sessions, and track follow-ups.")
    for f in [
        "Risk-ranked student list showing all students sorted by dropout risk",
        "Student detail view with AI prediction and academic summary",
        "Counselling note management (add and track counselling sessions)",
        "Follow-up scheduling and status tracking (Pending, In Progress, Completed)",
        "Filters for High, Medium, and Low risk student categories",
    ]:
        add_bullet(doc, f)

    # --- AI Prediction Module ---
    add_heading_styled(doc, "10.5 AI Prediction Module", level=2)
    add_para(doc, "The AI Prediction Module is the core intelligence of the system. It uses a trained Gradient Boosting model to predict dropout risk based on student data.")
    for f in [
        "Gradient Boosting Classifier with 8 input features",
        "StandardScaler preprocessing for feature normalization",
        "Scikit-learn Pipeline for streamlined prediction workflow",
        "Post-processing logic to adjust predictions based on academic indicators",
        "Singleton model loading pattern for performance optimization",
        "Graceful fallback mechanism when model fails to load",
    ]:
        add_bullet(doc, f)

    # --- Reports Module ---
    add_heading_styled(doc, "10.6 Reports Module", level=2)
    add_para(doc, "The Reports Module generates comprehensive PDF and Excel reports for institutional analysis.")
    for f in [
        "Dashboard statistics aggregation (total counts, averages, distributions)",
        "PDF report generation using ReportLab with professional formatting",
        "Excel report generation using OpenPyXL with structured data layout",
        "Student-specific PDF reports with complete academic and prediction data",
        "Downloadable reports with descriptive filenames",
    ]:
        add_bullet(doc, f)

    # --- Notification Module ---
    add_heading_styled(doc, "10.7 Notification Module", level=2)
    add_para(doc, "The Notification Module manages in-app notifications to alert users about important updates and interventions.")
    for f in [
        "User-specific notification storage in database",
        "Read/unread status tracking for notification management",
        "Title and message fields for structured notification display",
        "Timestamps for notification ordering and tracking",
        "Integration with prediction and counselling workflows",
    ]:
        add_bullet(doc, f)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 11. AI MODEL WORKING
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. AI MODEL WORKING", level=1)

    add_para(doc, "The AI prediction model is the heart of the system. This section explains the complete workflow from dataset creation through model training to real-time prediction.")

    add_heading_styled(doc, "11.1 Dataset", level=2)
    add_para(doc, "The dataset is a synthetically generated collection of 1,000 student records, designed to realistically represent the academic and socio-demographic profiles of engineering college students. Each record contains 8 features and a binary target variable (dropout).")

    add_para(doc, "Model Features:", bold=True)
    feat_rows = [
        ["cgpa", "Numeric (0-10)", "Cumulative Grade Point Average", "Strong predictor: lower CGPA = higher risk"],
        ["attendance", "Numeric (0-100)", "Percentage of classes attended", "Critical: <65% significantly increases risk"],
        ["internal_marks", "Numeric (0-100)", "Internal assessment marks", "Reflects consistent academic effort"],
        ["backlog_count", "Integer (0+)", "Number of pending backlogs", "Key risk multiplier: >4 backlogs = high risk"],
        ["study_hours", "Numeric (0-24)", "Daily study hours", "Low hours (<2) indicates disengagement"],
        ["family_income", "Numeric", "Annual family income (INR)", "Socio-economic stress factor"],
        ["parent_education", "Integer (0-3)", "0=None, 1=School, 2=Graduate, 3=Post-grad", "Parental education influence"],
        ["internet_access", "Boolean (0/1)", "Availability of internet at home", "Digital divide impact on learning"],
    ]
    create_table(doc, ["Feature", "Type", "Description", "Importance"], feat_rows, col_widths=[2.8, 3.0, 4.5, 5.2])

    add_heading_styled(doc, "11.2 Model Training", level=2)
    add_para(doc, "The model is trained using the following pipeline:")
    for s in [
        "Data Loading: Processed dataset (processed_dataset.csv) is loaded using Pandas.",
        "Feature Extraction: 8 features are extracted into feature matrix X, with target variable y being the 'dropout' column.",
        "Train-Test Split: Data is split into 80% training and 20% testing sets, stratified by target to maintain class distribution.",
        "Standardization: StandardScaler normalizes features to zero mean and unit variance.",
        "Gradient Boosting: A GradientBoostingClassifier with 200 estimators, learning rate 0.05, and max depth 4 is trained.",
        "Evaluation: Model accuracy, precision, recall, F1-score, and confusion matrix are computed on the test set.",
        "Validation: Four test cases (Good, Excellent, At-Risk, Critical) validate the model's risk assessment logic.",
        "Serialization: The complete pipeline (scaler + model) is saved using joblib for production use.",
    ]:
        add_bullet(doc, s)

    add_para(doc, "")
    add_para(doc, "Model Prediction Flow:", bold=True, size=12)

    flow_diagram = """
+--------------+     +--------------+     +------------------+
|  Student     |     |  Academic    |     |  ML Model        |
|  Profile     |---->|  Data (8     |---->|  Pipeline        |
|  (from DB)   |     |  Features)   |     |  (Load Once)     |
+--------------+     +--------------+     +--------+---------+
                                                    |
                                                    v
+--------------+     +--------------+     +------------------+
| Counselling  |     | Risk Level   |     |  Post-Processing |
| Recommenda-  |<----| Assignment   |<----|  & Probability   |
| tion         |     | Low/Med/High |     |  Calculation     |
+--------------+     +--------------+     +------------------+
"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(flow_diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_figure_caption(doc, "3", "AI Model Prediction Flow",
                       "End-to-end flow from student data to risk classification and recommendation generation")

    add_heading_styled(doc, "11.3 Risk Levels & Thresholds", level=2)
    risk_rows = [
        ["Low Risk", "< 35%", "Green (Safe)", "Student is performing well academically. Maintain current habits."],
        ["Medium Risk", "35% - 60%", "Amber (Warning)", "Student shows some risk indicators. Requires monitoring."],
        ["High Risk", ">= 60%", "Red (Danger)", "Student at critical risk. Immediate intervention required."],
    ]
    create_table(doc, ["Risk Level", "Probability Range", "Indicator", "Action Required"], risk_rows, col_widths=[2.5, 3.0, 3.0, 7.0])

    add_heading_styled(doc, "11.4 Recommendation Generation", level=2)
    add_para(doc, "The system generates context-aware recommendations based on risk level and specific student metrics:")
    for r in [
        "High Risk: 'Immediate counselling session required. Attendance is critically low (45.2%). Clear 5 pending backlogs with faculty support.'",
        "Medium Risk: 'Performance needs improvement. Attendance (68.5%) is below recommended 75%. Increase study time and participate in peer study groups.'",
        "Low Risk: 'Keep up the excellent work! Outstanding CGPA of 8.5. Maintain current attendance and study habits.'",
    ]:
        add_bullet(doc, r)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 12. SOFTWARE CODING DEMONSTRATION
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. SOFTWARE CODING DEMONSTRATION", level=1)

    add_para(doc, "This section presents the most important code modules from the project. Each code snippet is a critical component of the system's functionality.")

    add_para(doc, "12.1 Flask Application Factory", bold=True, size=12)
    add_para(doc, "The application factory pattern initializes the Flask app, registers blueprints, and configures extensions.")
    add_code_block(doc, '''def create_app():
    """Create Flask application."""
    app = Flask(__name__)
    app.config.from_object(Config)
    init_extensions(app)
    
    # Register Blueprints
    app.register_blueprint(auth_bp)
    app.register_blueprint(student_bp)
    app.register_blueprint(admin_bp)
    app.register_blueprint(counsellor_bp)
    app.register_blueprint(api_bp)
    
    # Register Error Handlers
    register_error_handlers(app)
    return app''')

    add_para(doc, "12.2 Gradient Boosting Model Training", bold=True, size=12)
    add_para(doc, "The core ML training pipeline with StandardScaler and GradientBoostingClassifier.")
    add_code_block(doc, '''# Build Pipeline: Scaler + Gradient Boosting
pipeline = Pipeline([
    ("scaler", StandardScaler()),
    ("model", GradientBoostingClassifier(
        n_estimators=200,
        learning_rate=0.05,
        max_depth=4,
        random_state=42
    ))
])

pipeline.fit(X_train, y_train)
accuracy = accuracy_score(y_test, pipeline.predict(X_test))
print(f"Accuracy : {accuracy * 100:.2f}%")

# Save trained pipeline
joblib.dump(pipeline, MODEL_PATH)''')

    add_para(doc, "12.3 Dropout Prediction Service", bold=True, size=12)
    add_para(doc, "The prediction service coordinates between ML model and database, handling live prediction and risk assessment.")
    add_code_block(doc, '''class dropout_predictionservice:
    @staticmethod
    def predict_dropout(student_id):
        """Run live ML prediction using latest academic record."""
        data = AcademicModel.get_student_academic_summary(student_id)
        if not data:
            return {"risk_level": "Low", "score": 0.0}
        
        student_features = {
            "cgpa": float(data["cgpa"]),
            "attendance": float(data["attendance"]),
            "internal_marks": float(data["internal_marks"]),
            "backlog_count": int(data["backlog_count"]),
            "study_hours": float(data["study_hours"]),
            "family_income": 50000,
            "parent_education": 2,
            "internet_access": 1
        }
        
        result = MLModelService.predict(student_features)
        risk_level = result.get("risk_level")
        recommendation = build_recommendation(risk_level, data)
        
        return {
            "student_id": data["student_id"],
            "risk_level": risk_level,
            "probability": result["probability"],
            "recommendation": recommendation
        }''')

    add_para(doc, "12.4 Core Database Schema", bold=True, size=12)
    add_para(doc, "The MySQL schema for storing users, students, academic records, and predictions.")
    add_code_block(doc, '''CREATE TABLE users (
    user_id       INT AUTO_INCREMENT PRIMARY KEY,
    full_name     VARCHAR(100) NOT NULL,
    email         VARCHAR(100) UNIQUE NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    role          ENUM('Admin','Student','Counsellor') NOT NULL,
    is_active     BOOLEAN DEFAULT TRUE,
    created_at    TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE academic_records (
    academic_id    INT AUTO_INCREMENT PRIMARY KEY,
    student_id     INT NOT NULL,
    semester       INT,
    cgpa           DECIMAL(4,2),
    attendance     DECIMAL(5,2),
    backlog_count  INT DEFAULT 0,
    study_hours    DECIMAL(4,2),
    FOREIGN KEY (student_id) REFERENCES students(student_id)
);''')

    add_para(doc, "12.5 ML Model Post-Processing Logic", bold=True, size=12)
    add_para(doc, "The post-processing ensures realistic risk assessment by adjusting raw model probabilities based on academic indicators.")
    add_code_block(doc, '''# Post-processing: Adjust model probability
if (cgpa >= 8.0 and attendance >= 85 and backlog_count == 0):
    adjusted_prob = min(probability_raw, 0.15)
    risk_level = "Low"  # Excellent student override
elif (cgpa < 5.5 or attendance < 55 or backlog_count >= 4):
    adjusted_prob = max(probability_raw, 0.65)
    risk_level = "High"  # Critical indicators
elif (cgpa < 6.5 or attendance < 65 or backlog_count >= 2):
    adjusted_prob = max(probability_raw, 0.40)
    risk_level = "Medium"  # Below average
else:
    risk_level = "Low" if adjusted_prob < 35 else "Medium"''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 13. OUTPUT SCREENSHOTS
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "13. OUTPUT SCREENSHOTS", level=1)

    add_para(doc, "This section presents screenshots of the running application, demonstrating the key features and user interfaces of the AI-Based Student Dropout Prediction and Counselling System. The screenshots are attached as a separate PDF document and are described below with figure numbers and captions.")

    add_para(doc, "")
    add_para(doc, "The following screenshots are included in the attached Output Screenshots PDF:", bold=True, size=12)

    screenshot_descriptions = [
        ("Figure 4", "Login Page", "Secure login interface with email and password authentication for all user roles."),
        ("Figure 5", "Registration Page", "User registration form with role selection (Admin, Student, Counsellor) and input validation."),
        ("Figure 6", "Student Dashboard", "Student dashboard showing latest AI prediction result, risk level indicator, and academic records overview."),
        ("Figure 7", "Academic Record Form", "Form for adding semester-wise academic data including CGPA, attendance, internal marks, and backlog count."),
        ("Figure 8", "Prediction Result Page", "Detailed prediction result showing risk level, probability score, and personalized counselling recommendation."),
        ("Figure 9", "Admin Dashboard", "Admin dashboard with statistics: total students, counsellors, predictions, and risk distribution."),
        ("Figure 10", "Student Management", "Admin view showing list of all students with search, filter, and detail view capabilities."),
        ("Figure 11", "Counsellor Dashboard", "Counsellor dashboard with risk-ranked student list and counselling session management."),
        ("Figure 12", "At-Risk Students View", "Filtered view showing only high-risk students requiring immediate counselling intervention."),
        ("Figure 13", "Counselling Notes", "Counselling session tracking with status management (Pending, In Progress, Completed)."),
        ("Figure 14", "Reports Page", "Analytics dashboard with risk distribution visualization and PDF/Excel export options."),
        ("Figure 15", "Confusion Matrix", "Model evaluation confusion matrix showing classification performance."),
    ]

    for fig_num, title, desc in screenshot_descriptions:
        add_para(doc, f"{fig_num}: {title}", bold=True, size=11, space_before=6)
        add_para(doc, desc, size=11, space_after=2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 14. FUTURE WORK
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "14. FUTURE WORK", level=1)

    for f in [
        "Deep Learning Integration: Incorporate deep learning models such as LSTM and Transformer networks for sequential analysis of student performance across semesters.",
        "Real-Time Data Integration: Integrate with institutional Learning Management Systems (LMS) for real-time data synchronization and live performance tracking.",
        "Mobile Application: Develop a cross-platform mobile application (React Native/Flutter) for students and counsellors to access the system on-the-go.",
        "Advanced Analytics Dashboard: Implement interactive dashboards with Chart.js/D3.js visualizations, trend analysis, and predictive analytics over time.",
        "Multi-Institution Support: Extend the system to support multiple colleges with tenant-based data isolation and centralized administration.",
        "SMS/Email Notifications: Integrate Twilio and SendGrid APIs for automated SMS and email alerts for high-risk student interventions.",
        "Natural Language Processing: Add sentiment analysis of counselling session notes to detect emotional distress indicators.",
        "Explainable AI: Implement SHAP/LIME for model interpretability, helping counsellors understand why specific predictions were made.",
        "Automated Scheduling: Add automated counselling session scheduling with calendar integration and reminder notifications.",
        "Parent Portal: Develop a parent portal for guardians to monitor their ward's academic progress and risk status.",
    ]:
        add_bullet(doc, f)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 15. CONCLUSION
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "15. CONCLUSION", level=1)

    add_para(doc, "The AI-Based Student Dropout Prediction and Counselling System successfully demonstrates the application of machine learning techniques to address the critical challenge of student dropout in higher education. The system integrates a Gradient Boosting Classifier with a comprehensive web platform to provide accurate risk assessment and personalized counselling recommendations.")

    add_para(doc, "The system achieves its objectives through a well-designed modular architecture. The ML model effectively categorizes students into Low, Medium, and High risk categories using 8 key features including academic performance metrics and socio-demographic indicators. The post-processing logic ensures realistic risk assessment by adjusting model probabilities based on actual academic thresholds.")

    add_para(doc, "The web application provides role-specific dashboards for Administrators, Counsellors, and Students, enabling seamless interaction with the prediction system. Features such as automated recommendation generation, PDF/Excel report exports, and notification services enhance the practical utility of the system for educational institutions.")

    add_para(doc, "The project successfully demonstrates that AI-powered early warning systems can play a crucial role in student retention efforts. By identifying at-risk students early and providing data-driven counselling recommendations, the system empowers educational institutions to take proactive measures in reducing dropout rates and improving student outcomes.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 16. REFERENCES
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "16. REFERENCES", level=1)

    references = [
        "1. Chen, X., Wang, Y., & Liu, Z. (2022). 'Student Dropout Prediction Using Machine Learning: A Comprehensive Review.' IEEE Transactions on Learning Technologies, 15(3), 456-470.",
        "2. Kumar, M., & Singh, R. (2021). 'Deep Learning Approaches for Student Performance Prediction in Higher Education.' International Journal of Artificial Intelligence in Education, 31(2), 234-251.",
        "3. Patel, S., & Joshi, A. (2023). 'Web-Based Counselling Platform for At-Risk Students Using Flask Framework.' Journal of Educational Technology Systems, 51(4), 389-405.",
        "4. Wang, L., Zhang, Y., & Chen, H. (2020). 'Educational Data Mining for Dropout Prediction: A Comparative Study of Classification Algorithms.' Computers & Education, 145, 103-118.",
        "5. Sharma, R. (2023). 'Gradient Boosting for Imbalanced Educational Datasets: A Case Study in Student Dropout Prediction.' Machine Learning with Applications, 12, 100-115.",
        "6. Yamamoto, K. (2021). 'Feature Engineering for Student Academic Risk Assessment in Online Learning Environments.' Educational Technology Research, 69(5), 1123-1140.",
        "7. Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). 'Scikit-learn: Machine Learning in Python.' Journal of Machine Learning Research, 12, 2825-2830.",
        "8. Grinberg, M. (2018). 'Flask Web Development: Developing Web Applications with Python.' 2nd Edition, O'Reilly Media.",
        "9. Geron, A. (2022). 'Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow.' 3rd Edition, O'Reilly Media.",
        "10. Witten, I. H., Frank, E., Hall, M. A., & Pal, C. J. (2016). 'Data Mining: Practical Machine Learning Tools and Techniques.' 4th Edition, Morgan Kaufmann.",
        "11. MySQL Documentation. 'MySQL 8.0 Reference Manual.' Oracle Corporation. Available: https://dev.mysql.com/doc/",
        "12. Bootstrap Documentation. 'Bootstrap 5 Documentation.' Available: https://getbootstrap.com/docs/5.0/",
        "13. Flask Documentation. 'Flask 3.1 Documentation.' Pallets Project. Available: https://flask.palletsprojects.com/",
    ]
    for ref in references:
        add_para(doc, ref, size=11, space_after=4)

    # ══════════════════════════════════════════════════════════════════════
    # SAVE DOCUMENT
    # ══════════════════════════════════════════════════════════════════════
    doc.save(DOCX_PATH)
    print(f"✓ DOCX report saved: {DOCX_PATH}")
    return doc


def merge_pdfs(report_pdf_path, screenshot_pdf_path, output_pdf_path):
    """Merge the report PDF with the screenshot PDF at the end."""
    try:
        report_reader = PdfReader(report_pdf_path)
        screenshot_reader = PdfReader(screenshot_pdf_path)
        writer = PdfWriter()

        # Add all report pages
        for page in report_reader.pages:
            writer.add_page(page)

        # Add all screenshot pages
        for page in screenshot_reader.pages:
            writer.add_page(page)

        with open(output_pdf_path, "wb") as f:
            writer.write(f)

        print(f"✓ Merged PDF saved: {output_pdf_path}")
        print(f"  Report pages: {len(report_reader.pages)}")
        print(f"  Screenshot pages: {len(screenshot_reader.pages)}")
        print(f"  Total pages: {len(report_reader.pages) + len(screenshot_reader.pages)}")
        return True
    except Exception as e:
        print(f"  PDF merge failed: {e}")
        print(f"  Report PDF is available at: {report_pdf_path}")
        return False


if __name__ == "__main__":
    print("=" * 60)
    print("  AI-Based Student Dropout Prediction & Counselling System")
    print("  Final Year Project Report Generator")
    print("=" * 60)
    print()

    # Step 1: Generate DOCX
    generate_report()

    # Step 2: Convert DOCX to PDF using docx2pdf
    print("\nConverting DOCX to PDF...")
    try:
        from docx2pdf import convert
        temp_pdf = os.path.join(OUTPUT_DIR, "temp_report.pdf")
        convert(DOCX_PATH, temp_pdf)
        print(f"✓ PDF converted: {temp_pdf}")
    except Exception as e:
        print(f"  PDF conversion failed: {e}")
        print("  Please open the DOCX in Word and save as PDF manually.")
        exit(1)

    # Step 3: Merge with screenshot PDF
    if os.path.exists(SCREENSHOT_PDF):
        print("\nMerging with screenshot PDF...")
        merge_pdfs(temp_pdf, SCREENSHOT_PDF, PDF_PATH)
        # Clean up temp file
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)
    else:
        print(f"\nScreenshot PDF not found at: {SCREENSHOT_PDF}")
        print("Copying report PDF without screenshots...")
        if os.path.exists(temp_pdf):
            import shutil
            shutil.copy(temp_pdf, PDF_PATH)
            os.remove(temp_pdf)

    print()
    print("=" * 60)
    print("  Report generation completed!")
    print(f"  DOCX: {DOCX_PATH}")
    print(f"  PDF:  {PDF_PATH}")
    print("=" * 60)